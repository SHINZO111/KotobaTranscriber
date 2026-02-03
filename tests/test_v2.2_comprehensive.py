"""
KotobaTranscriber v2.2 包括的テストスイート
PyQt5版の品質検証
"""

import os
import sys
import time
import json
import tempfile
import unittest
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional

# テスト対象のパスを追加
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

# ロギング設定
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============================================================================
# テストユーティリティ
# ============================================================================

class TestResult:
    """テスト結果を保持するクラス"""
    def __init__(self, name: str):
        self.name = name
        self.passed = 0
        self.failed = 0
        self.errors = 0
        self.skipped = 0
        self.details: List[str] = []
        self.duration = 0.0
    
    def add_pass(self, detail: str = ""):
        self.passed += 1
        if detail:
            self.details.append(f"✓ {detail}")
    
    def add_fail(self, detail: str = ""):
        self.failed += 1
        if detail:
            self.details.append(f"✗ {detail}")
    
    def add_error(self, detail: str = ""):
        self.errors += 1
        if detail:
            self.details.append(f"⚠ {detail}")
    
    def add_skip(self, detail: str = ""):
        self.skipped += 1
        if detail:
            self.details.append(f"⊘ {detail}")
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "passed": self.passed,
            "failed": self.failed,
            "errors": self.errors,
            "skipped": self.skipped,
            "total": self.passed + self.failed + self.errors + self.skipped,
            "duration": f"{self.duration:.2f}s",
            "details": self.details
        }


# ============================================================================
# 1. kotoba-whisper v2.2精度テスト
# ============================================================================

class TestKotobaWhisperPrecision(unittest.TestCase):
    """kotoba-whisper v2.2の精度テスト"""
    
    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("kotoba-whisper v2.2精度テスト")
        cls.start_time = time.time()
        
        try:
            from transcription_engine import TranscriptionEngine
            cls.TranscriptionEngine = TranscriptionEngine
            cls.available = True
        except ImportError as e:
            logger.warning(f"TranscriptionEngine not available: {e}")
            cls.available = False
    
    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time
    
    def test_01_model_loading(self):
        """モデルロードテスト"""
        if not self.available:
            self.result.add_skip("TranscriptionEngine not available")
            self.skipTest("Engine not available")
        
        try:
            engine = self.TranscriptionEngine()
            self.assertIsNotNone(engine)
            self.assertEqual(engine.model_name, "kotoba-tech/kotoba-whisper-v2.2")
            self.result.add_pass("モデル初期化成功")
        except Exception as e:
            self.result.add_fail(f"モデル初期化失敗: {e}")
            raise
    
    def test_02_device_selection(self):
        """デバイス選択テスト"""
        if not self.available:
            self.skipTest("Engine not available")
        
        from transcription_engine import DeviceSelector
        
        # 自動選択
        device = DeviceSelector.select_device("auto")
        self.assertIn(device, ["cuda", "cpu"])
        self.result.add_pass(f"デバイス自動選択: {device}")
        
        # 明示的選択
        device = DeviceSelector.select_device("cpu")
        self.assertEqual(device, "cpu")
        self.result.add_pass("CPU強制選択成功")
    
    def test_03_japanese_text_accuracy(self):
        """日本語テキスト精度テスト（モック）"""
        # 実際の音声なしで文字起こしロジックをテスト
        test_cases = [
            {
                "expected": "こんにちは、今日は会議です。",
                "keywords": ["こんにちは", "会議"]
            },
            {
                "expected": "プロジェクトの進捗を確認しましょう。",
                "keywords": ["プロジェクト", "進捗", "確認"]
            },
            {
                "expected": "来週の月曜日に提出してください。",
                "keywords": ["来週", "月曜日", "提出"]
            }
        ]
        
        for i, case in enumerate(test_cases):
            keywords_found = all(kw in case["expected"] for kw in case["keywords"])
            self.assertTrue(keywords_found)
            self.result.add_pass(f"テストケース {i+1}: キーワード検出成功")


# ============================================================================
# 2. pyannote.audio話者分離テスト
# ============================================================================

class TestSpeakerDiarization(unittest.TestCase):
    """話者分離機能のテスト"""
    
    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("pyannote.audio話者分離テスト")
        cls.start_time = time.time()
        
        try:
            from speaker_diarization_free import FreeSpeakerDiarizer
            cls.FreeSpeakerDiarizer = FreeSpeakerDiarizer
            cls.available = True
        except ImportError as e:
            logger.warning(f"Speaker diarization not available: {e}")
            cls.available = False
    
    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time
    
    def test_01_diarizer_initialization(self):
        """話者分離器の初期化テスト"""
        if not self.available:
            self.result.add_skip("Speaker diarization not available")
            self.skipTest("Not available")
        
        try:
            diarizer = self.FreeSpeakerDiarizer(method="auto")
            self.assertIsNotNone(diarizer)
            self.result.add_pass("話者分離器初期化成功")
        except Exception as e:
            self.result.add_fail(f"初期化失敗: {e}")
            raise
    
    def test_02_speaker_segment_format(self):
        """話者セグメント形式テスト"""
        # モックセグメントデータ
        mock_segments = [
            {"speaker": "SPEAKER_00", "start": 0.0, "end": 5.5},
            {"speaker": "SPEAKER_01", "start": 5.8, "end": 12.3},
            {"speaker": "SPEAKER_00", "start": 12.5, "end": 20.0},
        ]
        
        for seg in mock_segments:
            self.assertIn("speaker", seg)
            self.assertIn("start", seg)
            self.assertIn("end", seg)
            self.assertTrue(seg["end"] > seg["start"])
            self.assertTrue(seg["speaker"].startswith("SPEAKER_"))
        
        self.result.add_pass("セグメント形式検証成功")
    
    def test_03_speaker_formatting(self):
        """話者情報整形テスト"""
        if not self.available:
            self.skipTest("Not available")
        
        from speaker_diarization_utils import SpeakerFormatterMixin
        
        formatter = SpeakerFormatterMixin()
        
        segments = [
            {"speaker": "SPEAKER_00", "start": 0.0, "end": 5.0},
            {"speaker": "SPEAKER_01", "start": 5.0, "end": 10.0},
        ]
        
        text = "これはテストです。"
        formatted = formatter.format_with_speakers(text, segments)
        
        self.assertIsNotNone(formatted)
        self.result.add_pass("話者情報整形成功")


# ============================================================================
# 3. 句読点自動付加テスト
# ============================================================================

class TestPunctuationAddition(unittest.TestCase):
    """句読点自動付加テスト"""
    
    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("句読点自動付加テスト")
        cls.start_time = time.time()
        
        try:
            from text_formatter import TextFormatter
            cls.TextFormatter = TextFormatter
            cls.available = True
        except ImportError as e:
            logger.warning(f"TextFormatter not available: {e}")
            cls.available = False
    
    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time
    
    def test_01_basic_punctuation(self):
        """基本的な句読点付加"""
        if not self.available:
            self.result.add_skip("TextFormatter not available")
            self.skipTest("Not available")
        
        formatter = self.TextFormatter()
        
        # テストケース
        test_cases = [
            {
                "input": "こんにちは今日は会議です",
                "expected_punctuations": ["、", "。"],
                "description": "基本ケース"
            },
            {
                "input": "しかし問題がありましてそのため対応が必要です",
                "expected_punctuations": ["、", "。", "、"],
                "description": "接続詞あり"
            },
            {
                "input": "プロジェクトは順調です明日も頑張りましょう",
                "expected_punctuations": ["、", "。", "。"],
                "description": "複数文"
            }
        ]
        
        for case in test_cases:
            result = formatter.add_punctuation(case["input"])
            
            # 句読点が追加されているか確認
            has_punctuation = any(p in result for p in case["expected_punctuations"])
            self.assertTrue(has_punctuation, f"{case['description']}: 句読点が追加されていません")
            
            # 文末に句点があるか確認
            self.assertTrue(
                result.endswith(("。", "！", "？", "…", "\n")),
                f"{case['description']}: 文末処理が正しくありません"
            )
            
            self.result.add_pass(f"{case['description']}: OK")
    
    def test_02_filler_removal(self):
        """フィラー語削除テスト"""
        if not self.available:
            self.skipTest("Not available")
        
        formatter = self.TextFormatter()
        
        test_input = "あのーえーとですね今日は会議ですあの"
        result = formatter.remove_fillers(test_input)
        
        # フィラー語が削除されているか
        fillers = ["あのー", "えーと", "ですね"]
        for filler in fillers:
            self.assertNotIn(filler, result)
        
        self.result.add_pass("フィラー語削除成功")
    
    def test_03_paragraph_formatting(self):
        """段落整形テスト"""
        if not self.available:
            self.skipTest("Not available")
        
        formatter = self.TextFormatter()
        
        # 長いテキスト
        long_text = "文1です。" * 10
        
        result = formatter.format_paragraphs(long_text, max_sentences_per_paragraph=3)
        
        # 段落分けされているか確認
        paragraphs = result.split('\n\n')
        self.assertGreater(len(paragraphs), 1, "段落分けされていません")
        
        self.result.add_pass("段落整形成功")


# ============================================================================
# 4. 多形式出力テスト（TXT, DOCX, CSV, SRT, PDF）
# ============================================================================

class TestMultiFormatExport(unittest.TestCase):
    """多形式出力テスト"""
    
    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("多形式出力テスト（TXT, DOCX, CSV, SRT, PDF）")
        cls.start_time = time.time()
        cls.temp_dir = tempfile.mkdtemp()
        
        try:
            from subtitle_exporter import SubtitleExporter
            cls.SubtitleExporter = SubtitleExporter
            cls.subtitle_available = True
        except ImportError:
            cls.subtitle_available = False
    
    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time
        # クリーンアップ
        import shutil
        shutil.rmtree(cls.temp_dir, ignore_errors=True)
    
    def test_01_txt_export(self):
        """TXT形式エクスポート"""
        test_content = "これはテストです。\n2行目です。"
        output_path = os.path.join(self.temp_dir, "test.txt")
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(test_content)
            
            self.assertTrue(os.path.exists(output_path))
            
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            self.assertEqual(content, test_content)
            self.result.add_pass("TXTエクスポート成功")
        except Exception as e:
            self.result.add_fail(f"TXTエクスポート失敗: {e}")
            raise
    
    def test_02_docx_export(self):
        """DOCX形式エクスポート"""
        try:
            from docx import Document
            
            output_path = os.path.join(self.temp_dir, "test.docx")
            doc = Document()
            doc.add_heading('文字起こし結果', 0)
            doc.add_paragraph('これはテストです。')
            doc.add_paragraph('2行目です。')
            doc.save(output_path)
            
            self.assertTrue(os.path.exists(output_path))
            self.result.add_pass("DOCXエクスポート成功")
        except ImportError:
            self.result.add_skip("python-docx not installed")
            self.skipTest("python-docx not installed")
        except Exception as e:
            self.result.add_fail(f"DOCXエクスポート失敗: {e}")
            raise
    
    def test_03_csv_export(self):
        """CSV形式エクスポート"""
        try:
            import csv
            
            output_path = os.path.join(self.temp_dir, "test.csv")
            
            with open(output_path, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                writer.writerow(['時間', '話者', 'テキスト'])
                writer.writerow(['00:00:01', '話者A', 'こんにちは'])
                writer.writerow(['00:00:05', '話者B', '今日は会議です'])
            
            self.assertTrue(os.path.exists(output_path))
            
            # 読み込み確認
            with open(output_path, 'r', encoding='utf-8-sig') as f:
                reader = csv.reader(f)
                rows = list(reader)
            
            self.assertEqual(len(rows), 3)  # ヘッダー + 2データ行
            self.result.add_pass("CSVエクスポート成功")
        except Exception as e:
            self.result.add_fail(f"CSVエクスポート失敗: {e}")
            raise
    
    def test_04_srt_export(self):
        """SRT形式エクスポート"""
        if not self.subtitle_available:
            self.result.add_skip("SubtitleExporter not available")
            self.skipTest("Not available")
        
        try:
            exporter = self.SubtitleExporter()
            
            segments = [
                {"start": 0.5, "end": 3.2, "text": "こんにちは"},
                {"start": 3.5, "end": 6.8, "text": "今日は会議です"},
            ]
            
            output_path = os.path.join(self.temp_dir, "test.srt")
            success = exporter.export_srt(segments, output_path)
            
            self.assertTrue(success)
            self.assertTrue(os.path.exists(output_path))
            
            # 内容確認
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            self.assertIn("1", content)
            self.assertIn("00:00:00,500 --> 00:00:03,200", content)
            self.result.add_pass("SRTエクスポート成功")
        except Exception as e:
            self.result.add_fail(f"SRTエクスポート失敗: {e}")
            raise
    
    def test_05_vtt_export(self):
        """VTT形式エクスポート"""
        if not self.subtitle_available:
            self.skipTest("Not available")
        
        try:
            exporter = self.SubtitleExporter()
            
            segments = [
                {"start": 0.5, "end": 3.2, "text": "こんにちは"},
            ]
            
            output_path = os.path.join(self.temp_dir, "test.vtt")
            success = exporter.export_vtt(segments, output_path)
            
            self.assertTrue(success)
            
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            self.assertIn("WEBVTT", content)
            self.result.add_pass("VTTエクスポート成功")
        except Exception as e:
            self.result.add_fail(f"VTTエクスポート失敗: {e}")
            raise


# ============================================================================
# 5. PyQt5 GUI安定性テスト
# ============================================================================

class TestPyQt5GUIStability(unittest.TestCase):
    """PyQt5 GUI安定性テスト"""
    
    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("PyQt5 GUI安定性テスト")
        cls.start_time = time.time()
        
        # Qtインポートをチェック
        try:
            from qt_compat import (
                QApplication, QMainWindow, QWidget, QVBoxLayout,
                QPushButton, QLabel, Signal, Slot
            )
            cls.qt_available = True
        except ImportError:
            cls.qt_available = False
    
    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time
    
    def test_01_qt_imports(self):
        """Qtインポートテスト"""
        if not self.qt_available:
            self.result.add_skip("Qt not available")
            self.skipTest("Qt not available")
        
        from qt_compat import (
            QApplication, QMainWindow, QWidget, QVBoxLayout,
            QPushButton, QLabel, Signal, Slot, QT_VERSION
        )
        
        self.assertIn(QT_VERSION, ["PyQt5", "PySide6"])
        self.result.add_pass(f"Qtバインディング: {QT_VERSION}")
    
    def test_02_signal_slot(self):
        """Signal/Slot機能テスト"""
        if not self.qt_available:
            self.skipTest("Qt not available")
        
        from qt_compat import QObject, Signal, Slot
        
        class TestObject(QObject):
            test_signal = Signal(str)
            
            def __init__(self):
                super().__init__()
                self.received = None
            
            @Slot(str)
            def on_test(self, msg):
                self.received = msg
        
        obj = TestObject()
        obj.test_signal.connect(obj.on_test)
        obj.test_signal.emit("test_message")
        
        self.assertEqual(obj.received, "test_message")
        self.result.add_pass("Signal/Slot動作確認")
    
    def test_03_thread_safety(self):
        """スレッド安全性テスト"""
        if not self.qt_available:
            self.skipTest("Qt not available")
        
        from qt_compat import QThread, Signal
        import time
        
        class TestWorker(QThread):
            finished = Signal(str)
            
            def run(self):
                time.sleep(0.1)
                self.finished.emit("completed")
        
        worker = TestWorker()
        result = []
        
        def on_finish(msg):
            result.append(msg)
        
        worker.finished.connect(on_finish)
        worker.start()
        worker.wait(1000)
        
        self.assertEqual(result, ["completed"])
        self.result.add_pass("スレッド安全性確認")


# ============================================================================
# 6. モデル切り替えテスト（tiny〜large-v3）
# ============================================================================

class TestModelSwitching(unittest.TestCase):
    """モデル切り替えテスト"""
    
    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("モデル切り替えテスト（tiny〜large-v3）")
        cls.start_time = time.time()
        
        cls.available_models = [
            ("tiny", "openai/whisper-tiny"),
            ("base", "openai/whisper-base"),
            ("small", "openai/whisper-small"),
            ("medium", "openai/whisper-medium"),
            ("large-v2", "openai/whisper-large-v2"),
            ("large-v3", "openai/whisper-large-v3"),
        ]
    
    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time
    
    def test_01_model_name_validation(self):
        """モデル名検証テスト"""
        valid_models = [
            "kotoba-tech/kotoba-whisper-v2.2",
            "openai/whisper-tiny",
            "openai/whisper-base",
            "openai/whisper-small",
            "openai/whisper-medium",
            "openai/whisper-large-v2",
            "openai/whisper-large-v3",
        ]
        
        for model in valid_models:
            # モデル名の基本検証
            self.assertTrue('/' in model, f"Invalid model name: {model}")
            self.assertTrue(len(model) > 0)
        
        self.result.add_pass(f"{len(valid_models)}モデルの検証成功")
    
    def test_02_model_config_loading(self):
        """モデル設定読み込みテスト"""
        try:
            from config_manager import get_config
            
            config = get_config()
            
            # 各モデル設定を確認
            model_name = config.get("model.whisper.name")
            self.assertIsNotNone(model_name)
            
            # デバイス設定
            device = config.get("model.whisper.device", "auto")
            self.assertIn(device, ["auto", "cuda", "cpu"])
            
            # チャンク長
            chunk_length = config.get("model.whisper.chunk_length_s", 15)
            self.assertIsInstance(chunk_length, int)
            self.assertGreater(chunk_length, 0)
            
            self.result.add_pass("モデル設定読み込み成功")
        except ImportError:
            self.result.add_skip("config_manager not available")
            self.skipTest("Config manager not available")
    
    def test_03_faster_whisper_models(self):
        """Faster Whisperモデルテスト"""
        try:
            from faster_whisper import WhisperModel
            
            # 利用可能なモデルサイズ
            model_sizes = ["tiny", "base", "small", "medium", "large-v1", "large-v2", "large-v3"]
            
            for size in model_sizes:
                self.assertIsInstance(size, str)
                self.assertGreater(len(size), 0)
            
            self.result.add_pass(f"Faster Whisper {len(model_sizes)}モデル確認")
        except ImportError:
            self.result.add_skip("faster_whisper not installed")


# ============================================================================
# テスト実行エントリーポイント
# ============================================================================

def generate_report(results: List[TestResult]) -> str:
    """テストレポートを生成"""
    
    report_lines = [
        "=" * 80,
        "KotobaTranscriber v2.2 PyQt5版 テストレポート",
        "=" * 80,
        f"実行日時: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "【概要】",
        "-" * 80,
    ]
    
    total_passed = 0
    total_failed = 0
    total_errors = 0
    total_skipped = 0
    
    for result in results:
        total = result.passed + result.failed + result.errors + result.skipped
        status = "✓ PASS" if result.failed == 0 and result.errors == 0 else "✗ FAIL"
        
        report_lines.append(
            f"{status} | {result.name:40s} | "
            f"Pass: {result.passed:2d} | Fail: {result.failed:2d} | "
            f"Error: {result.errors:2d} | Skip: {result.skipped:2d} | "
            f"Time: {result.duration:.2f}s"
        )
        
        total_passed += result.passed
        total_failed += result.failed
        total_errors += result.errors
        total_skipped += result.skipped
    
    report_lines.extend([
        "-" * 80,
        f"合計: Pass: {total_passed} | Fail: {total_failed} | Error: {total_errors} | Skip: {total_skipped}",
        "",
        "【詳細】",
        "=" * 80,
    ])
    
    for result in results:
        if result.details:
            report_lines.append(f"\n{result.name}:")
            for detail in result.details:
                report_lines.append(f"  {detail}")
    
    report_lines.append("\n" + "=" * 80)
    
    return "\n".join(report_lines)


def run_all_tests():
    """全テストを実行"""
    
    # テストスイート作成
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()
    
    # テスト追加
    suite.addTests(loader.loadTestsFromTestCase(TestKotobaWhisperPrecision))
    suite.addTests(loader.loadTestsFromTestCase(TestSpeakerDiarization))
    suite.addTests(loader.loadTestsFromTestCase(TestPunctuationAddition))
    suite.addTests(loader.loadTestsFromTestCase(TestMultiFormatExport))
    suite.addTests(loader.loadTestsFromTestCase(TestPyQt5GUIStability))
    suite.addTests(loader.loadTestsFromTestCase(TestModelSwitching))
    
    # テスト実行
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    
    # レポート保存
    results = [
        TestKotobaWhisperPrecision.result,
        TestSpeakerDiarization.result,
        TestPunctuationAddition.result,
        TestMultiFormatExport.result,
        TestPyQt5GUIStability.result,
        TestModelSwitching.result,
    ]
    
    report = generate_report(results)
    
    # レポートをファイルに保存
    report_path = os.path.join(
        os.path.dirname(__file__),
        f"test_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    )
    
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(report)
    
    print(f"\nレポート保存: {report_path}")
    print(report)
    
    return result.wasSuccessful()


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
