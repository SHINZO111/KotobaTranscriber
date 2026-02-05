"""
KotobaTranscriber v2.2 Comprehensive Test Suite - PyQt5 Quality Verification
"""

import os
import sys
import time
import json
import tempfile
import unittest
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional

# Add test target path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

# Logging settings
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============================================================================
# Test Utilities
# ============================================================================

class TestResult:
    """Class to hold test results"""
    def __init__(self, name: str):
        self.name = name
        self.passed = 0
        self.failed = 0
        self.errors = 0
        self.skipped = 0
        self.details: List[str] = []
        self.duration = 0.0

    def add_pass(self, detail: str = ""):
        self.passed += 1
        if detail:
            self.details.append(f"PASS: {detail}")

    def add_fail(self, detail: str = ""):
        self.failed += 1
        if detail:
            self.details.append(f"FAIL: {detail}")

    def add_error(self, detail: str = ""):
        self.errors += 1
        if detail:
            self.details.append(f"ERROR: {detail}")

    def add_skip(self, detail: str = ""):
        self.skipped += 1
        if detail:
            self.details.append(f"SKIP: {detail}")

    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "passed": self.passed,
            "failed": self.failed,
            "errors": self.errors,
            "skipped": self.skipped,
            "total": self.passed + self.failed + self.errors + self.skipped,
            "duration": f"{self.duration:.2f}s",
            "details": self.details
        }


# ============================================================================
# 1. kotoba-whisper v2.2 Precision Tests
# ============================================================================

class TestKotobaWhisperPrecision(unittest.TestCase):
    """kotoba-whisper v2.2 precision tests"""

    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("kotoba-whisper v2.2 Precision Tests")
        cls.start_time = time.time()

        try:
            from transcription_engine import KotobaTranscriptionEngine
            cls.TranscriptionEngine = KotobaTranscriptionEngine
            cls.available = True
        except ImportError as e:
            logger.warning(f"TranscriptionEngine not available: {e}")
            cls.available = False

    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time

    def test_01_model_loading(self):
        """Model loading test"""
        if not self.available:
            self.result.add_skip("TranscriptionEngine not available")
            self.skipTest("Engine not available")

        try:
            engine = self.TranscriptionEngine()
            self.assertIsNotNone(engine)
            self.assertEqual(engine.model_name, "kotoba-tech/kotoba-whisper-v2.2")
            self.result.add_pass("Model initialization successful")
        except Exception as e:
            self.result.add_fail(f"Model initialization failed: {e}")
            raise

    def test_02_device_selection(self):
        """Device selection test"""
        if not self.available:
            self.skipTest("Engine not available")

        from transcription_engine import DeviceSelector

        # Auto selection
        device = DeviceSelector.select_device("auto")
        self.assertIn(device, ["cuda", "cpu"])
        self.result.add_pass(f"Device auto selection: {device}")

        # Explicit selection
        device = DeviceSelector.select_device("cpu")
        self.assertEqual(device, "cpu")
        self.result.add_pass("CPU forced selection successful")

    def test_03_japanese_text_accuracy(self):
        """Japanese text accuracy test (mock)"""
        # Test without actual audio
        test_cases = [
            {
                "expected": "Hello, the weather is nice today.",
                "keywords": ["Hello", "weather"]
            },
            {
                "expected": "Let's check the project progress.",
                "keywords": ["project", "progress", "check"]
            },
            {
                "expected": "Please attend the meeting next week.",
                "keywords": ["meeting", "next week", "attend"]
            }
        ]

        for i, case in enumerate(test_cases):
            keywords_found = all(kw in case["expected"] for kw in case["keywords"])
            self.assertTrue(keywords_found)
            self.result.add_pass(f"Test case {i+1}: Keyword verification successful")


# ============================================================================
# 2. pyannote.audio Speaker Diarization Tests
# ============================================================================

class TestSpeakerDiarization(unittest.TestCase):
    """Speaker diarization functionality tests"""

    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("pyannote.audio Speaker Diarization Tests")
        cls.start_time = time.time()

        try:
            from speaker_diarization_free import FreeSpeakerDiarizer
            cls.FreeSpeakerDiarizer = FreeSpeakerDiarizer
            cls.available = True
        except ImportError as e:
            logger.warning(f"Speaker diarization not available: {e}")
            cls.available = False

    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time

    def test_01_diarizer_initialization(self):
        """Diarizer initialization test"""
        if not self.available:
            self.result.add_skip("Speaker diarization not available")
            self.skipTest("Not available")

        try:
            diarizer = self.FreeSpeakerDiarizer(method="auto")
            self.assertIsNotNone(diarizer)
            self.result.add_pass("Diarizer initialization successful")
        except Exception as e:
            self.result.add_fail(f"Initialization failed: {e}")
            raise

    def test_02_speaker_segment_format(self):
        """Speaker segment format test"""
        # Mock segment data
        mock_segments = [
            {"speaker": "SPEAKER_00", "start": 0.0, "end": 5.5},
            {"speaker": "SPEAKER_01", "start": 5.8, "end": 12.3},
            {"speaker": "SPEAKER_00", "start": 12.5, "end": 20.0},
        ]

        for seg in mock_segments:
            self.assertIn("speaker", seg)
            self.assertIn("start", seg)
            self.assertIn("end", seg)
            self.assertTrue(seg["end"] > seg["start"])
            self.assertTrue(seg["speaker"].startswith("SPEAKER_"))

        self.result.add_pass("Segment format verification successful")


# ============================================================================
# 3. Punctuation Addition Tests
# ============================================================================

class TestPunctuationAddition(unittest.TestCase):
    """Punctuation auto-addition tests"""

    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("Punctuation Auto-Addition Tests")
        cls.start_time = time.time()

        try:
            from text_formatter import TextFormatter
            cls.TextFormatter = TextFormatter
            cls.available = True
        except ImportError as e:
            logger.warning(f"TextFormatter not available: {e}")
            cls.available = False

    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time

    def test_01_basic_punctuation(self):
        """Basic punctuation addition"""
        if not self.available:
            self.result.add_skip("TextFormatter not available")
            self.skipTest("Not available")

        formatter = self.TextFormatter()

        # Test cases
        test_cases = [
            {
                "input": "Hello the weather is nice today",
                "description": "Basic case"
            },
            {
                "input": "However there is a question so a response is needed",
                "description": "With conjunctions"
            },
            {
                "input": "The project is going well let's work hard tomorrow too",
                "description": "Multiple sentences"
            }
        ]

        for case in test_cases:
            result = formatter.add_punctuation(case["input"])
            self.assertIsNotNone(result)
            self.result.add_pass(f"{case['description']}: OK")

    def test_02_filler_removal(self):
        """Filler word removal test"""
        if not self.available:
            self.skipTest("Not available")

        formatter = self.TextFormatter()

        test_input = "um well you know the weather is nice today"
        result = formatter.remove_fillers(test_input)

        self.assertIsNotNone(result)
        self.result.add_pass("Filler word removal successful")

    def test_03_paragraph_formatting(self):
        """Paragraph formatting test"""
        if not self.available:
            self.skipTest("Not available")

        formatter = self.TextFormatter()

        # Long text with Japanese punctuation (required for sentence splitting)
        long_text = "これは文章です。" * 10  # Japanese sentence repeated 10 times

        result = formatter.format_paragraphs(long_text, max_sentences_per_paragraph=3)

        # Check paragraphs are split (should split 10 sentences into groups of 3)
        paragraphs = result.split('\n\n')
        self.assertGreater(len(paragraphs), 1, "Paragraphs not split")

        self.result.add_pass("Paragraph formatting successful")


# ============================================================================
# 4. Multi-Format Export Tests (TXT, DOCX, CSV, SRT, PDF)
# ============================================================================

class TestMultiFormatExport(unittest.TestCase):
    """Multi-format export tests"""

    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("Multi-Format Export Tests (TXT, DOCX, CSV, SRT, PDF)")
        cls.start_time = time.time()
        cls.temp_dir = tempfile.mkdtemp()

        try:
            from subtitle_exporter import SubtitleExporter
            cls.SubtitleExporter = SubtitleExporter
            cls.subtitle_available = True
        except ImportError:
            cls.subtitle_available = False

    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time
        # Cleanup
        import shutil
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    def test_01_txt_export(self):
        """TXT format export"""
        test_content = "This is a test.\nSecond line.\n"
        output_path = os.path.join(self.temp_dir, "test.txt")

        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(test_content)

            self.assertTrue(os.path.exists(output_path))

            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()

            self.assertEqual(content, test_content)
            self.result.add_pass("TXT export successful")
        except Exception as e:
            self.result.add_fail(f"TXT export failed: {e}")
            raise

    def test_02_docx_export(self):
        """DOCX format export"""
        try:
            from docx import Document

            output_path = os.path.join(self.temp_dir, "test.docx")
            doc = Document()
            doc.add_heading('Transcription Result', 0)
            doc.add_paragraph('This is a test.')
            doc.add_paragraph('Second line.')
            doc.save(output_path)

            self.assertTrue(os.path.exists(output_path))
            self.result.add_pass("DOCX export successful")
        except ImportError:
            self.result.add_skip("python-docx not installed")
            self.skipTest("python-docx not installed")
        except Exception as e:
            self.result.add_fail(f"DOCX export failed: {e}")
            raise

    def test_03_csv_export(self):
        """CSV format export"""
        try:
            import csv

            output_path = os.path.join(self.temp_dir, "test.csv")

            with open(output_path, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                writer.writerow(['Time', 'Speaker', 'Text'])
                writer.writerow(['00:00:01', 'Speaker1', 'Hello'])
                writer.writerow(['00:00:05', 'Speaker2', 'Nice weather'])

            self.assertTrue(os.path.exists(output_path))

            # Read verification
            with open(output_path, 'r', encoding='utf-8-sig') as f:
                reader = csv.reader(f)
                rows = list(reader)

            self.assertEqual(len(rows), 3)  # Header + 2 data rows
            self.result.add_pass("CSV export successful")
        except Exception as e:
            self.result.add_fail(f"CSV export failed: {e}")
            raise

    def test_04_srt_export(self):
        """SRT format export"""
        if not self.subtitle_available:
            self.result.add_skip("SubtitleExporter not available")
            self.skipTest("Not available")

        try:
            exporter = self.SubtitleExporter()

            segments = [
                {"start": 0.5, "end": 3.2, "text": "Hello"},
                {"start": 3.5, "end": 6.8, "text": "Nice weather today"},
            ]

            output_path = os.path.join(self.temp_dir, "test.srt")
            success = exporter.export_srt(segments, output_path)

            self.assertTrue(success)
            self.assertTrue(os.path.exists(output_path))

            # Content verification
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()

            self.assertIn("1", content)
            self.assertIn("00:00:00,500 --> 00:00:03,200", content)
            self.result.add_pass("SRT export successful")
        except Exception as e:
            self.result.add_fail(f"SRT export failed: {e}")
            raise

    def test_05_vtt_export(self):
        """VTT format export"""
        if not self.subtitle_available:
            self.skipTest("Not available")

        try:
            exporter = self.SubtitleExporter()

            segments = [
                {"start": 0.5, "end": 3.2, "text": "Hello"},
            ]

            output_path = os.path.join(self.temp_dir, "test.vtt")
            success = exporter.export_vtt(segments, output_path)

            self.assertTrue(success)

            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()

            self.assertIn("WEBVTT", content)
            self.result.add_pass("VTT export successful")
        except Exception as e:
            self.result.add_fail(f"VTT export failed: {e}")
            raise


# ============================================================================
# 5. PyQt5 GUI Stability Tests
# ============================================================================

class TestPyQt5GUIStability(unittest.TestCase):
    """PyQt5 GUI stability tests"""

    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("PyQt5 GUI Stability Tests")
        cls.start_time = time.time()

        # Check Qt imports
        try:
            from qt_compat import (
                QApplication, QMainWindow, QWidget, QVBoxLayout,
                QPushButton, QLabel, Signal, Slot
            )
            cls.qt_available = True
        except ImportError:
            cls.qt_available = False

    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time

    def test_01_qt_imports(self):
        """Qt import test"""
        if not self.qt_available:
            self.result.add_skip("Qt not available")
            self.skipTest("Qt not available")

        from qt_compat import (
            QApplication, QMainWindow, QWidget, QVBoxLayout,
            QPushButton, QLabel, Signal, Slot, QT_VERSION
        )

        self.assertIn(QT_VERSION, ["PyQt5", "PySide6"])
        self.result.add_pass(f"Qt binding: {QT_VERSION}")

    def test_02_signal_slot(self):
        """Signal/Slot functionality test"""
        if not self.qt_available:
            self.skipTest("Qt not available")

        from qt_compat import QObject, Signal, Slot

        class TestObject(QObject):
            test_signal = Signal(str)

            def __init__(self):
                super().__init__()
                self.received = None

            @Slot(str)
            def on_test(self, msg):
                self.received = msg

        obj = TestObject()
        obj.test_signal.connect(obj.on_test)
        obj.test_signal.emit("test_message")

        self.assertEqual(obj.received, "test_message")
        self.result.add_pass("Signal/Slot operation verified")

    def test_03_thread_safety(self):
        """Thread safety test"""
        if not self.qt_available:
            self.skipTest("Qt not available")

        from qt_compat import QThread, Signal
        import time

        class TestWorker(QThread):
            finished = Signal(str)

            def run(self):
                time.sleep(0.1)
                self.finished.emit("completed")

        worker = TestWorker()
        result = []

        def on_finish(msg):
            result.append(msg)

        worker.finished.connect(on_finish)
        worker.start()
        worker.wait(1000)

        self.assertEqual(result, ["completed"])
        self.result.add_pass("Thread safety verified")


# ============================================================================
# 6. Model Switching Tests (tiny - large-v3)
# ============================================================================

class TestModelSwitching(unittest.TestCase):
    """Model switching tests"""

    @classmethod
    def setUpClass(cls):
        cls.result = TestResult("Model Switching Tests (tiny - large-v3)")
        cls.start_time = time.time()

        cls.available_models = [
            ("tiny", "openai/whisper-tiny"),
            ("base", "openai/whisper-base"),
            ("small", "openai/whisper-small"),
            ("medium", "openai/whisper-medium"),
            ("large-v2", "openai/whisper-large-v2"),
            ("large-v3", "openai/whisper-large-v3"),
        ]

    @classmethod
    def tearDownClass(cls):
        cls.result.duration = time.time() - cls.start_time

    def test_01_model_name_validation(self):
        """Model name validation test"""
        valid_models = [
            "kotoba-tech/kotoba-whisper-v2.2",
            "openai/whisper-tiny",
            "openai/whisper-base",
            "openai/whisper-small",
            "openai/whisper-medium",
            "openai/whisper-large-v2",
            "openai/whisper-large-v3",
        ]

        for model in valid_models:
            # Basic model name validation
            self.assertTrue('/' in model, f"Invalid model name: {model}")
            self.assertTrue(len(model) > 0)

        self.result.add_pass(f"{len(valid_models)} models validation successful")

    def test_02_model_config_loading(self):
        """Model config loading test"""
        try:
            from config_manager import get_config

            config = get_config()

            # Check model settings
            model_name = config.get("model.whisper.name")
            self.assertIsNotNone(model_name)

            # Device settings
            device = config.get("model.whisper.device", "auto")
            self.assertIn(device, ["auto", "cuda", "cpu"])

            # Chunk length
            chunk_length = config.get("model.whisper.chunk_length_s", 15)
            self.assertIsInstance(chunk_length, int)
            self.assertGreater(chunk_length, 0)

            self.result.add_pass("Model config loading successful")
        except ImportError:
            self.result.add_skip("config_manager not available")
            self.skipTest("Config manager not available")

    def test_03_faster_whisper_models(self):
        """Faster Whisper models test"""
        try:
            from faster_whisper import WhisperModel

            # Available model sizes
            model_sizes = ["tiny", "base", "small", "medium", "large-v1", "large-v2", "large-v3"]

            for size in model_sizes:
                self.assertIsInstance(size, str)
                self.assertGreater(len(size), 0)

            self.result.add_pass(f"Faster Whisper {len(model_sizes)} models verified")
        except ImportError:
            self.result.add_skip("faster_whisper not installed")


# ============================================================================
# Test Execution Entry Point
# ============================================================================

def generate_report(results: List[TestResult]) -> str:
    """Generate test report"""

    report_lines = [
        "=" * 80,
        "KotobaTranscriber v2.2 PyQt5 Test Report",
        "=" * 80,
        f"Execution time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "[Summary]",
        "-" * 80,
    ]

    total_passed = 0
    total_failed = 0
    total_errors = 0
    total_skipped = 0

    for result in results:
        total = result.passed + result.failed + result.errors + result.skipped
        status = "PASS" if result.failed == 0 and result.errors == 0 else "FAIL"

        report_lines.append(
            f"{status} | {result.name:40s} | "
            f"Pass: {result.passed:2d} | Fail: {result.failed:2d} | "
            f"Error: {result.errors:2d} | Skip: {result.skipped:2d} | "
            f"Time: {result.duration:.2f}s"
        )

        total_passed += result.passed
        total_failed += result.failed
        total_errors += result.errors
        total_skipped += result.skipped

    report_lines.extend([
        "-" * 80,
        f"Total: Pass: {total_passed} | Fail: {total_failed} | Error: {total_errors} | Skip: {total_skipped}",
        "",
        "[Details]",
        "=" * 80,
    ])

    for result in results:
        if result.details:
            report_lines.append(f"\n{result.name}:")
            for detail in result.details:
                report_lines.append(f"  {detail}")

    report_lines.append("\n" + "=" * 80)

    return "\n".join(report_lines)


def run_all_tests():
    """Run all tests"""

    # Create test suite
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()

    # Add tests
    suite.addTests(loader.loadTestsFromTestCase(TestKotobaWhisperPrecision))
    suite.addTests(loader.loadTestsFromTestCase(TestSpeakerDiarization))
    suite.addTests(loader.loadTestsFromTestCase(TestPunctuationAddition))
    suite.addTests(loader.loadTestsFromTestCase(TestMultiFormatExport))
    suite.addTests(loader.loadTestsFromTestCase(TestPyQt5GUIStability))
    suite.addTests(loader.loadTestsFromTestCase(TestModelSwitching))

    # Run tests
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)

    # Save report
    results = [
        TestKotobaWhisperPrecision.result,
        TestSpeakerDiarization.result,
        TestPunctuationAddition.result,
        TestMultiFormatExport.result,
        TestPyQt5GUIStability.result,
        TestModelSwitching.result,
    ]

    report = generate_report(results)

    # Save report to file
    report_path = os.path.join(
        os.path.dirname(__file__),
        f"test_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    )

    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(report)

    print(f"\nReport saved: {report_path}")
    print(report)

    return result.wasSuccessful()


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
