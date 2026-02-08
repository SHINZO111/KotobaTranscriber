"""
強化エクスポートモジュール
Excel形式(.xlsx)とWord形式(.docx)の議事録出力対応
"""

import logging
import threading
from pathlib import Path
from typing import List, Dict, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import io

logger = logging.getLogger(__name__)


@dataclass
class ExportOptions:
    """エクスポートオプション"""
    include_timestamps: bool = True
    include_speaker_labels: bool = True
    merge_short_segments: bool = False
    min_segment_duration: float = 1.0
    format_type: str = "meeting"  # meeting, transcript, subtitle
    template_path: Optional[str] = None
    company_name: str = "AGEC株式会社"
    project_name: str = ""


class ExcelExporter:
    """Excel形式エクスポーター"""

    def __init__(self):
        self.has_openpyxl = self._check_openpyxl()

    def _check_openpyxl(self) -> bool:
        """openpyxlがインストールされているかチェック"""
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.utils import get_column_letter
            return True
        except ImportError:
            logger.warning("openpyxl not installed. Excel export will be unavailable.")
            return False

    def export_transcription(
        self,
        segments: List[Dict],
        output_path: str,
        options: Optional[ExportOptions] = None,
    ) -> bool:
        """
        書き起こしをExcel形式でエクスポート

        Args:
            segments: 書き起こしセグメント
            output_path: 出力ファイルパス
            options: エクスポートオプション

        Returns:
            成功したかどうか
        """
        if not self.has_openpyxl:
            logger.error("openpyxl is required for Excel export")
            return False

        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "書き起こし"

            # ヘッダー行
            headers = ["No.", "開始時間", "終了時間", "話者", "テキスト"]
            ws.append(headers)

            # ヘッダースタイル
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center")

            # データ行
            for i, segment in enumerate(segments, 1):
                start_time = self._format_time(segment.get("start", 0))
                end_time = self._format_time(segment.get("end", 0))
                speaker = segment.get("speaker", "Unknown")
                text = segment.get("text", "")

                ws.append([i, start_time, end_time, speaker, text])

            # 列幅調整
            ws.column_dimensions["A"].width = 6
            ws.column_dimensions["B"].width = 12
            ws.column_dimensions["C"].width = 12
            ws.column_dimensions["D"].width = 15
            ws.column_dimensions["E"].width = 60

            # 枠線
            thin_border = Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            )
            for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=5):
                for cell in row:
                    cell.border = thin_border

            # 自動折り返し
            for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=5, max_col=5):
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical="top")

            wb.save(output_path)
            logger.info(f"Excel transcription exported to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export Excel transcription: {e}")
            return False

    def export_meeting_minutes(
        self,
        minutes_data: Dict,
        output_path: str,
        options: Optional[ExportOptions] = None,
    ) -> bool:
        """
        議事録をExcel形式でエクスポート（AGEC社内テンプレート対応）

        Args:
            minutes_data: 議事録データ
            output_path: 出力ファイルパス
            options: エクスポートオプション

        Returns:
            成功したかどうか
        """
        if not self.has_openpyxl:
            logger.error("openpyxl is required for Excel export")
            return False

        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "議事録"

            options = options or ExportOptions()
            company = options.company_name

            # タイトル
            ws.merge_cells("A1:E1")
            title_cell = ws["A1"]
            title_cell.value = f"{company} 議事録"
            title_cell.font = Font(bold=True, size=16)
            title_cell.alignment = Alignment(horizontal="center", vertical="center")
            ws.row_dimensions[1].height = 30

            # 基本情報
            row = 3
            ws[f"A{row}"] = "会議名"
            ws[f"B{row}"] = minutes_data.get("title", "")
            ws.merge_cells(f"B{row}:E{row}")
            row += 1

            ws[f"A{row}"] = "日時"
            ws[f"B{row}"] = minutes_data.get("date", "")
            ws[f"C{row}"] = "場所"
            ws[f"D{row}"] = minutes_data.get("location", "")
            ws.merge_cells(f"D{row}:E{row}")
            row += 1

            # 出席者
            ws[f"A{row}"] = "出席者"
            attendees = minutes_data.get("attendees", [])
            ws[f"B{row}"] = ", ".join(attendees) if attendees else ""
            ws.merge_cells(f"B{row}:E{row}")
            row += 2

            # 議題
            if minutes_data.get("agenda"):
                ws[f"A{row}"] = "議題"
                ws[f"A{row}"].font = Font(bold=True)
                row += 1
                for agenda_item in minutes_data["agenda"]:
                    ws[f"B{row}"] = f"• {agenda_item}"
                    ws.merge_cells(f"B{row}:E{row}")
                    row += 1
                row += 1

            # 決定事項
            if minutes_data.get("decisions"):
                ws[f"A{row}"] = "決定事項"
                ws[f"A{row}"].font = Font(bold=True)
                ws[f"A{row}"].fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
                row += 1
                for i, decision in enumerate(minutes_data["decisions"], 1):
                    ws[f"A{row}"] = f"{i}."
                    ws[f"B{row}"] = decision
                    ws.merge_cells(f"B{row}:E{row}")
                    row += 1
                row += 1

            # 確認事項
            if minutes_data.get("confirmations"):
                ws[f"A{row}"] = "確認事項"
                ws[f"A{row}"].font = Font(bold=True)
                ws[f"A{row}"].fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
                row += 1
                for i, confirmation in enumerate(minutes_data["confirmations"], 1):
                    ws[f"A{row}"] = f"{i}."
                    ws[f"B{row}"] = confirmation
                    ws.merge_cells(f"B{row}:E{row}")
                    row += 1
                row += 1

            # アクションアイテム
            if minutes_data.get("action_items"):
                ws[f"A{row}"] = "アクションアイテム"
                ws[f"A{row}"].font = Font(bold=True)
                ws[f"A{row}"].fill = PatternFill(start_color="D9EAD3", end_color="D9EAD3", fill_type="solid")
                ws[f"B{row}"] = "内容"
                ws[f"C{row}"] = "担当"
                ws[f"D{row}"] = "期限"
                ws[f"E{row}"] = "優先度"
                for cell in ws[row]:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="D9EAD3", end_color="D9EAD3", fill_type="solid")
                row += 1

                for item in minutes_data["action_items"]:
                    ws[f"A{row}"] = "□"
                    ws[f"B{row}"] = item.get("description", "")
                    ws[f"C{row}"] = item.get("assignee", "")
                    ws[f"D{row}"] = item.get("due_date", "")
                    ws[f"E{row}"] = item.get("priority", "中")
                    row += 1
                row += 1

            # 議事内容
            if minutes_data.get("statements"):
                ws[f"A{row}"] = "議事内容"
                ws[f"A{row}"].font = Font(bold=True)
                row += 1

                current_speaker = None
                for stmt in minutes_data["statements"]:
                    speaker = stmt.get("speaker", "Unknown")
                    text = stmt.get("text", "")
                    stmt_type = stmt.get("statement_type", "")

                    if speaker != current_speaker:
                        ws[f"A{row}"] = speaker
                        ws[f"A{row}"].font = Font(bold=True, color="4472C4")
                        current_speaker = speaker
                    else:
                        ws[f"A{row}"] = ""

                    prefix = ""
                    if stmt_type == "決定事項":
                        prefix = "【決定】"
                    elif stmt_type == "アクションアイテム":
                        prefix = "【TODO】"
                    elif stmt_type == "確認事項":
                        prefix = "【確認】"

                    ws[f"B{row}"] = f"{prefix}{text}"
                    ws.merge_cells(f"B{row}:E{row}")
                    row += 1

            # 次回会議
            if minutes_data.get("next_meeting"):
                row += 1
                ws[f"A{row}"] = "次回会議"
                ws[f"A{row}"].font = Font(bold=True)
                ws[f"B{row}"] = minutes_data["next_meeting"]
                ws.merge_cells(f"B{row}:E{row}")

            # 列幅調整
            ws.column_dimensions["A"].width = 12
            ws.column_dimensions["B"].width = 50
            ws.column_dimensions["C"].width = 15
            ws.column_dimensions["D"].width = 15
            ws.column_dimensions["E"].width = 12

            # 枠線
            thin_border = Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            )
            for row_cells in ws.iter_rows():
                for cell in row_cells:
                    if cell.value:
                        cell.border = thin_border

            wb.save(output_path)
            logger.info(f"Excel meeting minutes exported to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export Excel meeting minutes: {e}")
            return False

    def _format_time(self, seconds: float) -> str:
        """秒数を時:分:秒形式に変換"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        if hours > 0:
            return f"{hours}:{minutes:02d}:{secs:02d}"
        return f"{minutes:02d}:{secs:02d}"


class WordExporter:
    """Word形式エクスポーター"""

    def __init__(self):
        self.has_python_docx = self._check_python_docx()

    def _check_python_docx(self) -> bool:
        """python-docxがインストールされているかチェック"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            return True
        except ImportError:
            logger.warning("python-docx not installed. Word export will be unavailable.")
            return False

    def export_transcription(
        self,
        segments: List[Dict],
        output_path: str,
        options: Optional[ExportOptions] = None,
    ) -> bool:
        """
        書き起こしをWord形式でエクスポート

        Args:
            segments: 書き起こしセグメント
            output_path: 出力ファイルパス
            options: エクスポートオプション

        Returns:
            成功したかどうか
        """
        if not self.has_python_docx:
            logger.error("python-docx is required for Word export")
            return False

        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            doc = Document()

            # タイトル
            title = doc.add_heading("書き起こしテキスト", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 日付
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.add_run(f"作成日: {datetime.now().strftime('%Y年%m月%d日')}")

            doc.add_paragraph()  # 空行

            # 内容
            current_speaker = None
            for segment in segments:
                speaker = segment.get("speaker", "Unknown")
                text = segment.get("text", "")
                start = segment.get("start", 0)

                if speaker != current_speaker:
                    # 話者見出し
                    speaker_para = doc.add_paragraph()
                    speaker_run = speaker_para.add_run(f"【{speaker}】")
                    speaker_run.bold = True
                    speaker_run.font.size = Pt(11)
                    speaker_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    current_speaker = speaker

                # テキスト（タイムスタンプ付き）
                time_str = self._format_time(start)
                text_para = doc.add_paragraph(style="List Bullet")
                text_para.add_run(f"[{time_str}] ").bold = True
                text_para.add_run(text)

            doc.save(output_path)
            logger.info(f"Word transcription exported to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export Word transcription: {e}")
            return False

    def export_meeting_minutes(
        self,
        minutes_data: Dict,
        output_path: str,
        options: Optional[ExportOptions] = None,
    ) -> bool:
        """
        議事録をWord形式でエクスポート（AGEC社内テンプレート対応）

        Args:
            minutes_data: 議事録データ
            output_path: 出力ファイルパス
            options: エクスポートオプション

        Returns:
            成功したかどか
        """
        if not self.has_python_docx:
            logger.error("python-docx is required for Word export")
            return False

        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            doc = Document()

            options = options or ExportOptions()
            company = options.company_name

            # タイトル
            title = doc.add_heading(f"{company} 議事録", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 基本情報テーブル
            table = doc.add_table(rows=3, cols=4)
            table.style = "Light Grid Accent 1"

            # 会議名
            table.rows[0].cells[0].text = "会議名"
            table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
            table.rows[0].cells[1].merge(table.rows[0].cells[3])
            table.rows[0].cells[1].text = minutes_data.get("title", "")

            # 日時・場所
            table.rows[1].cells[0].text = "日時"
            table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
            table.rows[1].cells[1].text = minutes_data.get("date", "")
            table.rows[1].cells[2].text = "場所"
            table.rows[1].cells[2].paragraphs[0].runs[0].bold = True
            table.rows[1].cells[3].text = minutes_data.get("location", "")

            # 出席者
            table.rows[2].cells[0].text = "出席者"
            table.rows[2].cells[0].paragraphs[0].runs[0].bold = True
            table.rows[2].cells[1].merge(table.rows[2].cells[3])
            attendees = minutes_data.get("attendees", [])
            table.rows[2].cells[1].text = ", ".join(attendees) if attendees else ""

            doc.add_paragraph()  # 空行

            # 議題
            if minutes_data.get("agenda"):
                doc.add_heading("議題", level=1)
                for agenda_item in minutes_data["agenda"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(agenda_item)
                doc.add_paragraph()

            # 決定事項
            if minutes_data.get("decisions"):
                doc.add_heading("決定事項", level=1)
                for i, decision in enumerate(minutes_data["decisions"], 1):
                    p = doc.add_paragraph(style="List Number")
                    run = p.add_run(decision)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)  # 紺色
                doc.add_paragraph()

            # 確認事項
            if minutes_data.get("confirmations"):
                doc.add_heading("確認事項", level=1)
                for i, confirmation in enumerate(minutes_data["confirmations"], 1):
                    p = doc.add_paragraph(style="List Number")
                    run = p.add_run(confirmation)
                    run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)  # オレンジ色
                doc.add_paragraph()

            # アクションアイテム
            if minutes_data.get("action_items"):
                doc.add_heading("アクションアイテム", level=1)
                for item in minutes_data["action_items"]:
                    p = doc.add_paragraph()
                    p.add_run("☐ ").font.size = Pt(14)
                    p.add_run(item.get("description", "")).bold = True

                    # 詳細情報
                    details = []
                    if item.get("assignee"):
                        details.append(f"担当: {item['assignee']}")
                    if item.get("due_date"):
                        details.append(f"期限: {item['due_date']}")
                    if item.get("priority"):
                        details.append(f"優先度: {item['priority']}")

                    if details:
                        detail_p = doc.add_paragraph()
                        detail_p.paragraph_format.left_indent = Inches(0.3)
                        detail_run = detail_p.add_run(", ".join(details))
                        detail_run.font.size = Pt(9)
                        detail_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                doc.add_paragraph()

            # 議事内容
            if minutes_data.get("statements"):
                doc.add_heading("議事内容", level=1)

                current_speaker = None
                for stmt in minutes_data["statements"]:
                    speaker = stmt.get("speaker", "Unknown")
                    text = stmt.get("text", "")
                    stmt_type = stmt.get("statement_type", "")

                    if speaker != current_speaker:
                        p = doc.add_paragraph()
                        run = p.add_run(f"【{speaker}】")
                        run.bold = True
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                        current_speaker = speaker

                    prefix = ""
                    if stmt_type == "決定事項":
                        prefix = "【決定】"
                    elif stmt_type == "アクションアイテム":
                        prefix = "【TODO】"
                    elif stmt_type == "確認事項":
                        prefix = "【確認】"

                    p = doc.add_paragraph(style="List Bullet")
                    if prefix:
                        run = p.add_run(prefix)
                        run.bold = True
                    p.add_run(text)

            # 次回会議
            if minutes_data.get("next_meeting"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("次回会議: ").bold = True
                p.add_run(minutes_data["next_meeting"])

            # 備考
            if minutes_data.get("notes"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("備考:\n").bold = True
                p.add_run(minutes_data["notes"])

            doc.save(output_path)
            logger.info(f"Word meeting minutes exported to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export Word meeting minutes: {e}")
            return False

    def _format_time(self, seconds: float) -> str:
        """秒数を時:分:秒形式に変換"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        if hours > 0:
            return f"{hours}:{minutes:02d}:{secs:02d}"
        return f"{minutes:02d}:{secs:02d}"


class EnhancedExporter:
    """強化エクスポートクラス（統合インターフェース）"""

    def __init__(self):
        self.excel_exporter = ExcelExporter()
        self.word_exporter = WordExporter()

    def export(
        self,
        data: Dict,
        output_path: str,
        format_type: str = "excel",
        data_type: str = "meeting_minutes",
        options: Optional[ExportOptions] = None,
    ) -> bool:
        """
        データを指定形式でエクスポート

        Args:
            data: エクスポートするデータ
            output_path: 出力ファイルパス
            format_type: 形式 ("excel", "word", "txt", "md")
            data_type: データタイプ ("meeting_minutes", "transcription")
            options: エクスポートオプション

        Returns:
            成功したかどうか
        """
        options = options or ExportOptions()

        if format_type in ["excel", "xlsx"]:
            if data_type == "meeting_minutes":
                return self.excel_exporter.export_meeting_minutes(data, output_path, options)
            else:
                return self.excel_exporter.export_transcription(data.get("segments", []), output_path, options)

        elif format_type in ["word", "docx"]:
            if data_type == "meeting_minutes":
                return self.word_exporter.export_meeting_minutes(data, output_path, options)
            else:
                return self.word_exporter.export_transcription(data.get("segments", []), output_path, options)

        elif format_type == "txt":
            return self._export_text(data, output_path, data_type)

        elif format_type in ["md", "markdown"]:
            return self._export_markdown(data, output_path, data_type)

        else:
            logger.error(f"Unsupported format: {format_type}")
            return False

    def _export_text(self, data: Dict, output_path: str, data_type: str) -> bool:
        """テキスト形式でエクスポート"""
        try:
            if data_type == "meeting_minutes" and "text_format" in data:
                content = data["text_format"]
            else:
                # 簡易テキスト出力
                lines = []
                for segment in data.get("segments", []):
                    speaker = segment.get("speaker", "Unknown")
                    text = segment.get("text", "")
                    lines.append(f"[{speaker}] {text}")
                content = "\n".join(lines)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)

            logger.info(f"Text exported to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export text: {e}")
            return False

    def _export_markdown(self, data: Dict, output_path: str, data_type: str) -> bool:
        """Markdown形式でエクスポート"""
        try:
            if data_type == "meeting_minutes" and "markdown_format" in data:
                content = data["markdown_format"]
            else:
                # 簡易Markdown出力
                lines = ["# 書き起こし", ""]
                for segment in data.get("segments", []):
                    speaker = segment.get("speaker", "Unknown")
                    text = segment.get("text", "")
                    lines.append(f"- **{speaker}**: {text}")
                content = "\n".join(lines)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)

            logger.info(f"Markdown exported to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export markdown: {e}")
            return False

    def get_available_formats(self) -> List[str]:
        """
        利用可能なフォーマットリストを取得

        Returns:
            フォーマットリスト
        """
        formats = ["txt", "md", "markdown"]
        if self.excel_exporter.has_openpyxl:
            formats.extend(["excel", "xlsx"])
        if self.word_exporter.has_python_docx:
            formats.extend(["word", "docx"])
        return formats


# グローバルインスタンス
_enhanced_exporter = None
_enhanced_exporter_lock = threading.Lock()


def get_enhanced_exporter() -> EnhancedExporter:
    """
    強化エクスポーターのシングルトンインスタンスを取得

    Returns:
        EnhancedExporterインスタンス
    """
    global _enhanced_exporter
    if _enhanced_exporter is None:
        with _enhanced_exporter_lock:
            if _enhanced_exporter is None:
                _enhanced_exporter = EnhancedExporter()
    return _enhanced_exporter


if __name__ == "__main__":
    # テスト用コード
    logging.basicConfig(level=logging.INFO)

    exporter = get_enhanced_exporter()

    # テストデータ
    test_minutes = {
        "title": "新規店舗開発会議",
        "date": "2026年2月3日 14:00-15:30",
        "location": "会議室A",
        "attendees": ["田中（PM）", "佐藤（設計）", "山田（工事）", "鈴木（コスト）"],
        "agenda": ["設計進捗確認", "施工業者選定", "予算確認"],
        "statements": [
            {"speaker": "田中", "text": "本日の会議を始めます。", "statement_type": ""},
            {"speaker": "佐藤", "text": "設計図は80%完了しています。", "statement_type": "報告"},
            {"speaker": "山田", "text": "外壁材はタイルに決定しました。", "statement_type": "決定事項"},
            {"speaker": "田中", "text": "施工業者との調整をお願いします。", "statement_type": "アクションアイテム"},
        ],
        "decisions": ["外壁材はタイルを採用", "工事開始は3月1日"],
        "confirmations": ["内装費は予定通り", "空調設備の仕様"],
        "action_items": [
            {"description": "施工業者との調整", "assignee": "佐藤", "due_date": "2/10", "priority": "高"},
            {"description": "予算内訳の確認", "assignee": "鈴木", "due_date": "2/7", "priority": "中"},
        ],
        "next_meeting": "2月10日（月）14:00～",
    }

    # Excel出力テスト
    if exporter.excel_exporter.has_openpyxl:
        output_path = "test_meeting_minutes.xlsx"
        success = exporter.export(test_minutes, output_path, "excel", "meeting_minutes")
        print(f"Excel export: {'成功' if success else '失敗'}")

    # Word出力テスト
    if exporter.word_exporter.has_python_docx:
        output_path = "test_meeting_minutes.docx"
        success = exporter.export(test_minutes, output_path, "word", "meeting_minutes")
        print(f"Word export: {'成功' if success else '失敗'}")

    print(f"\n利用可能なフォーマット: {exporter.get_available_formats()}")
