"""
Wordエクスポートモジュール
議事録・書き起こしのWord形式出力（python-docx使用）
"""

import logging
from pathlib import Path
from typing import List, Dict, Optional, Any
from dataclasses import dataclass
from datetime import datetime

logger = logging.getLogger(__name__)


@dataclass
class ExportOptions:
    """エクスポートオプション"""
    include_timestamps: bool = True
    include_speaker_labels: bool = True
    format_type: str = "meeting"
    template_path: Optional[str] = None
    company_name: str = "AGEC株式会社"
    project_name: str = ""


class WordExporter:
    """Word形式エクスポーター（.docx）"""

    def __init__(self):
        self.has_python_docx = self._check_python_docx()

    def _check_python_docx(self) -> bool:
        """python-docxがインストールされているかチェック"""
        try:
            from docx import Document
            return True
        except ImportError:
            logger.warning("python-docx not installed. Word export unavailable.")
            return False

    def export_transcription(
        self,
        segments: List[Dict],
        output_path: str,
        options: Optional[ExportOptions] = None,
    ) -> bool:
        """
        書き起こしをWord形式でエクスポート

        Args:
            segments: 書き起こしセグメント
            output_path: 出力ファイルパス
            options: エクスポートオプション

        Returns:
            成功したかどうか
        """
        if not self.has_python_docx:
            logger.error("python-docx is required for Word export")
            return False

        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            doc = Document()

            # タイトル
            title = doc.add_heading("書き起こしテキスト", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 日付
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.add_run(f"作成日: {datetime.now().strftime('%Y年%m月%d日')}")

            doc.add_paragraph()

            # 内容
            current_speaker = None
            for segment in segments:
                speaker = segment.get("speaker", "Unknown")
                text = segment.get("text", "")
                start = segment.get("start", 0)

                if speaker != current_speaker:
                    speaker_para = doc.add_paragraph()
                    speaker_run = speaker_para.add_run(f"【{speaker}】")
                    speaker_run.bold = True
                    speaker_run.font.size = Pt(11)
                    speaker_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    current_speaker = speaker

                time_str = self._format_time(start)
                text_para = doc.add_paragraph(style="List Bullet")
                text_para.add_run(f"[{time_str}] ").bold = True
                text_para.add_run(text)

            doc.save(output_path)
            logger.info(f"Word transcription exported to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export Word transcription: {e}")
            return False

    def export_meeting_minutes(
        self,
        minutes_data: Dict,
        output_path: str,
        options: Optional[ExportOptions] = None,
    ) -> bool:
        """
        議事録をWord形式でエクスポート（AGEC社内テンプレート対応）

        Args:
            minutes_data: 議事録データ
            output_path: 出力ファイルパス
            options: エクスポートオプション

        Returns:
            成功したかどうか
        """
        if not self.has_python_docx:
            logger.error("python-docx is required for Word export")
            return False

        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            doc = Document()

            options = options or ExportOptions()
            company = options.company_name

            # タイトル
            title = doc.add_heading(f"{company} 議事録", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 基本情報テーブル
            table = doc.add_table(rows=3, cols=4)
            table.style = "Light Grid Accent 1"

            # 会議名
            table.rows[0].cells[0].text = "会議名"
            table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
            table.rows[0].cells[1].merge(table.rows[0].cells[3])
            table.rows[0].cells[1].text = minutes_data.get("title", "")

            # 日時・場所
            table.rows[1].cells[0].text = "日時"
            table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
            table.rows[1].cells[1].text = minutes_data.get("date", "")
            table.rows[1].cells[2].text = "場所"
            table.rows[1].cells[2].paragraphs[0].runs[0].bold = True
            table.rows[1].cells[3].text = minutes_data.get("location", "")

            # 出席者
            table.rows[2].cells[0].text = "出席者"
            table.rows[2].cells[0].paragraphs[0].runs[0].bold = True
            table.rows[2].cells[1].merge(table.rows[2].cells[3])
            attendees = minutes_data.get("attendees", [])
            table.rows[2].cells[1].text = ", ".join(attendees) if attendees else ""

            doc.add_paragraph()

            # 議題
            if minutes_data.get("agenda"):
                doc.add_heading("議題", level=1)
                for agenda_item in minutes_data["agenda"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(agenda_item)
                doc.add_paragraph()

            # 決定事項
            if minutes_data.get("decisions"):
                doc.add_heading("決定事項", level=1)
                for i, decision in enumerate(minutes_data["decisions"], 1):
                    p = doc.add_paragraph(style="List Number")
                    run = p.add_run(decision)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
                doc.add_paragraph()

            # 確認事項
            if minutes_data.get("confirmations"):
                doc.add_heading("確認事項", level=1)
                for i, confirmation in enumerate(minutes_data["confirmations"], 1):
                    p = doc.add_paragraph(style="List Number")
                    run = p.add_run(confirmation)
                    run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
                doc.add_paragraph()

            # アクションアイテム
            if minutes_data.get("action_items"):
                doc.add_heading("アクションアイテム", level=1)
                for item in minutes_data["action_items"]:
                    p = doc.add_paragraph()
                    p.add_run("☐ ").font.size = Pt(14)
                    p.add_run(item.get("description", "")).bold = True

                    details = []
                    if item.get("assignee"):
                        details.append(f"担当: {item['assignee']}")
                    if item.get("due_date"):
                        details.append(f"期限: {item['due_date']}")
                    if item.get("priority"):
                        details.append(f"優先度: {item['priority']}")

                    if details:
                        detail_p = doc.add_paragraph()
                        detail_p.paragraph_format.left_indent = Inches(0.3)
                        detail_run = detail_p.add_run(", ".join(details))
                        detail_run.font.size = Pt(9)
                        detail_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                doc.add_paragraph()

            # 議事内容
            if minutes_data.get("statements"):
                doc.add_heading("議事内容", level=1)

                current_speaker = None
                for stmt in minutes_data["statements"]:
                    speaker = stmt.get("speaker", "Unknown")
                    text = stmt.get("text", "")
                    stmt_type = stmt.get("statement_type", "")

                    if speaker != current_speaker:
                        p = doc.add_paragraph()
                        run = p.add_run(f"【{speaker}】")
                        run.bold = True
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                        current_speaker = speaker

                    prefix = ""
                    if stmt_type == "決定事項":
                        prefix = "【決定】"
                    elif stmt_type == "アクションアイテム":
                        prefix = "【TODO】"
                    elif stmt_type == "確認事項":
                        prefix = "【確認】"

                    p = doc.add_paragraph(style="List Bullet")
                    if prefix:
                        run = p.add_run(prefix)
                        run.bold = True
                    p.add_run(text)

            # 次回会議
            if minutes_data.get("next_meeting"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("次回会議: ").bold = True
                p.add_run(minutes_data["next_meeting"])

            # 備考
            if minutes_data.get("notes"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("備考:\n").bold = True
                p.add_run(minutes_data["notes"])

            doc.save(output_path)
            logger.info(f"Word meeting minutes exported to {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to export Word meeting minutes: {e}")
            return False

    def _format_time(self, seconds: float) -> str:
        """秒数を時:分:秒形式に変換"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        if hours > 0:
            return f"{hours}:{minutes:02d}:{secs:02d}"
        return f"{minutes:02d}:{secs:02d}"


# グローバルインスタンス
_word_exporter = None


def get_word_exporter() -> WordExporter:
    """Wordエクスポーターのシングルトンインスタンスを取得"""
    global _word_exporter
    if _word_exporter is None:
        _word_exporter = WordExporter()
    return _word_exporter
