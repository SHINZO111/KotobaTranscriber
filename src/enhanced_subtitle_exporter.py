"""
enhanced_subtitle_exporter.py - 強化字幕エクスポートモジュール

SRT/VTT/JSON/DOCX形式へのエクスポートをサポート
"""

import html
import json
import re
from datetime import datetime
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod
from pathlib import Path
import logging
from time_utils import format_time_srt, format_time_vtt
from export.common import atomic_write_text, validate_export_path, validate_segments

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class SubtitleFormatter(ABC):
    """字幕フォーマッタ抽象基底クラス"""
    
    @abstractmethod
    def format_time(self, seconds: float) -> str:
        """時間をフォーマット"""
        pass
    
    @abstractmethod
    def format_segment(self, segment: Dict, index: int) -> str:
        """セグメントをフォーマット"""
        pass
    
    @abstractmethod
    def generate_header(self) -> str:
        """ヘッダーを生成"""
        pass
    
    def format_segments(self, segments: List[Dict]) -> str:
        """全セグメントをフォーマット"""
        lines = [self.generate_header()]
        
        for i, segment in enumerate(segments, 1):
            formatted = self.format_segment(segment, i)
            if formatted:
                lines.append(formatted)
        
        return '\n'.join(lines)


class SRTFormatter(SubtitleFormatter):
    """SRTフォーマッタ"""
    
    def format_time(self, seconds: float) -> str:
        """秒数をSRT時間形式に変換 (HH:MM:SS,mmm)"""
        return format_time_srt(seconds)
    
    def format_segment(self, segment: Dict, index: int) -> str:
        """SRTセグメントをフォーマット"""
        start = self.format_time(segment.get('start', 0))
        end = self.format_time(segment.get('end', 0))
        text = segment.get('text', '').strip()
        
        if not text:
            return ""
        
        # 話者情報を追加
        speaker = segment.get('speaker')
        if speaker:
            text = f"[{html.escape(speaker)}] {text}"

        return f"{index}\n{start} --> {end}\n{text}\n"
    
    def generate_header(self) -> str:
        """SRTヘッダー（空）"""
        return ""


class VTTFormatter(SubtitleFormatter):
    """WebVTTフォーマッタ"""
    
    def format_time(self, seconds: float) -> str:
        """秒数をVTT時間形式に変換 (HH:MM:SS.mmm)"""
        return format_time_vtt(seconds)
    
    def format_segment(self, segment: Dict, index: int) -> str:
        """VTTセグメントをフォーマット"""
        start = self.format_time(segment.get('start', 0))
        end = self.format_time(segment.get('end', 0))
        text = segment.get('text', '').strip()
        
        if not text:
            return ""
        
        # 話者情報を追加
        speaker = segment.get('speaker')
        if speaker:
            text = f"<v {html.escape(speaker)}>{html.escape(text)}</v>"
        
        return f"{start} --> {end}\n{text}\n"
    
    def generate_header(self) -> str:
        """VTTヘッダー"""
        return "WEBVTT\n\n"


class JSONFormatter:
    """JSONフォーマッタ"""
    
    def format(
        self, 
        segments: List[Dict], 
        metadata: Optional[Dict] = None
    ) -> str:
        """
        JSON形式で出力
        
        Args:
            segments: セグメントリスト
            metadata: メタデータ
            
        Returns:
            JSON文字列
        """
        output = {
            'version': '1.0',
            'format': 'kotoba-transcription',
            'generated_at': datetime.now().isoformat(),
            'segment_count': len(segments),
            'segments': segments,
            'metadata': metadata or {}
        }
        return json.dumps(output, ensure_ascii=False, indent=2)


class TXTFormatter:
    """プレーンテキストフォーマッタ"""
    
    def format(
        self, 
        segments: List[Dict],
        include_timestamps: bool = True,
        include_speakers: bool = True
    ) -> str:
        """
        プレーンテキストで出力
        
        Args:
            segments: セグメントリスト
            include_timestamps: タイムスタンプを含める
            include_speakers: 話者情報を含める
            
        Returns:
            テキスト文字列
        """
        lines = []
        
        for segment in segments:
            parts = []
            
            if include_timestamps:
                start = segment.get('start', 0)
                parts.append(f"[{start:06.2f}s]")
            
            if include_speakers and segment.get('speaker'):
                parts.append(f"[{segment['speaker']}]")
            
            parts.append(segment.get('text', '').strip())
            
            if parts:
                lines.append(' '.join(parts))
        
        return '\n'.join(lines)


class DOCXFormatter:
    """Wordドキュメントフォーマッタ"""
    
    def format(
        self,
        segments: List[Dict],
        output_path: str,
        title: str = "文字起こし結果",
        include_speakers: bool = True
    ) -> bool:
        """
        Wordドキュメントで出力
        
        Args:
            segments: セグメントリスト
            output_path: 出力パス
            title: ドキュメントタイトル
            include_speakers: 話者情報を含める
            
        Returns:
            True: 成功
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return False
        
        try:
            doc = Document()
            
            # タイトル
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # セグメントを追加
            current_speaker = None
            
            for segment in segments:
                speaker = segment.get('speaker')
                text = segment.get('text', '').strip()
                start = segment.get('start', 0)
                
                if not text:
                    continue
                
                # 話者変更時に見出し
                if include_speakers and speaker and speaker != current_speaker:
                    doc.add_heading(speaker, level=2)
                    current_speaker = speaker
                
                # 段落追加
                para = doc.add_paragraph()
                
                # タイムスタンプ（小字）
                run = para.add_run(f"[{start:06.2f}s] ")
                run.font.size = Pt(8)
                run.font.color.rgb = None  # グレーに設定可能
                
                # テキスト
                para.add_run(text)
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            return False


class EnhancedSubtitleExporter:
    """
    強化字幕エクスポーター
    
    複数フォーマットに対応した字幕エクスポート
    """
    
    FORMATTERS = {
        'srt': SRTFormatter,
        'vtt': VTTFormatter,
    }
    
    def __init__(self):
        self.json_formatter = JSONFormatter()
        self.txt_formatter = TXTFormatter()
        self.docx_formatter = DOCXFormatter()
    
    def export(
        self,
        segments: List[Dict],
        output_path: str,
        format_type: str,
        options: Optional[Dict] = None
    ) -> bool:
        """
        エクスポート実行
        
        Args:
            segments: セグメントリスト
            output_path: 出力パス
            format_type: フォーマット（srt, vtt, json, txt, docx）
            options: 追加オプション
            
        Returns:
            True: 成功
        """
        options = options or {}
        format_type = format_type.lower()

        try:
            validate_export_path(output_path)
            validate_segments(segments)
            if format_type in self.FORMATTERS:
                return self._export_subtitle(segments, output_path, format_type)
            
            elif format_type == 'json':
                content = self.json_formatter.format(
                    segments,
                    options.get('metadata')
                )
                atomic_write_text(output_path, content)
                return True

            elif format_type == 'txt':
                content = self.txt_formatter.format(
                    segments,
                    include_timestamps=options.get('include_timestamps', True),
                    include_speakers=options.get('include_speakers', True)
                )
                atomic_write_text(output_path, content)
                return True
            
            elif format_type == 'docx':
                return self.docx_formatter.format(
                    segments,
                    output_path,
                    title=options.get('title', '文字起こし結果'),
                    include_speakers=options.get('include_speakers', True)
                )
            
            else:
                raise ValueError(f"Unsupported format: {format_type}")
                
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return False
    
    def _export_subtitle(
        self, 
        segments: List[Dict], 
        output_path: str, 
        format_type: str
    ) -> bool:
        """字幕ファイルをエクスポート"""
        formatter_class = self.FORMATTERS[format_type]
        formatter = formatter_class()
        content = formatter.format_segments(segments)
        atomic_write_text(output_path, content)
        return True
    
    def export_auto(
        self,
        segments: List[Dict],
        base_path: str,
        formats: List[str] = None,
        options: Optional[Dict] = None
    ) -> Dict[str, bool]:
        """
        複数フォーマットで自動エクスポート
        
        Args:
            segments: セグメントリスト
            base_path: 基本パス（拡張子なし）
            formats: 出力フォーマットリスト
            options: 追加オプション
            
        Returns:
            {フォーマット: 成功/失敗}
        """
        if formats is None:
            formats = ['srt', 'txt']
        
        results = {}
        base = Path(base_path)
        
        for fmt in formats:
            ext = f".{fmt}"
            if fmt == 'txt':
                ext = "_字幕.txt"
            elif fmt == 'docx':
                ext = ".docx"
            
            output_path = base.with_suffix(ext)
            results[fmt] = self.export(
                segments, 
                str(output_path), 
                fmt, 
                options
            )
        
        return results
    
    def merge_short_segments(
        self,
        segments: List[Dict],
        min_duration: float = 1.0,
        max_chars: int = 40
    ) -> List[Dict]:
        """
        短いセグメントをマージ
        
        Args:
            segments: 元のセグメントリスト
            min_duration: 最小表示時間（秒）
            max_chars: 最大文字数
            
        Returns:
            マージ済みセグメントリスト
        """
        if not segments:
            return []
        
        merged = []
        current = None
        
        for segment in segments:
            seg_start = segment.get('start', 0)
            seg_end = segment.get('end', 0)
            seg_text = segment.get('text', '').strip()

            if current is None:
                current = {
                    'start': seg_start,
                    'end': seg_end,
                    'text': seg_text,
                    'speaker': segment.get('speaker')
                }
                continue

            duration = seg_end - current['start']
            combined_text = current['text'] + ' ' + seg_text

            # マージ条件
            same_speaker = segment.get('speaker') == current.get('speaker')
            can_merge = (
                duration < min_duration and
                len(combined_text) <= max_chars and
                same_speaker
            )

            if can_merge:
                current['end'] = seg_end
                current['text'] = combined_text
            else:
                merged.append(current)
                current = {
                    'start': seg_start,
                    'end': seg_end,
                    'text': seg_text,
                    'speaker': segment.get('speaker')
                }
        
        if current:
            merged.append(current)
        
        return merged
    
    def split_long_segments(
        self,
        segments: List[Dict],
        max_chars: int = 40,
        max_duration: float = 5.0
    ) -> List[Dict]:
        """
        長いセグメントを分割
        
        Args:
            segments: 元のセグメントリスト
            max_chars: 最大文字数
            max_duration: 最大表示時間（秒）
            
        Returns:
            分割済みセグメントリスト
        """
        result = []
        
        for segment in segments:
            text = segment.get('text', '').strip()
            start = segment.get('start', 0)
            end = segment.get('end', 0)
            duration = end - start
            
            # 分割不要チェック
            if len(text) <= max_chars and duration <= max_duration:
                result.append(segment)
                continue
            
            # 文で分割
            sentences = re.split(r'([。！？\.\!\?])', text)
            sentences = [s for s in sentences if s]
            
            # 文を結合して適切な長さに
            current_text = ""
            current_start = start
            time_per_char = duration / len(text) if text else 0
            
            for i, sentence in enumerate(sentences):
                if not current_text:
                    current_text = sentence
                elif len(current_text) + len(sentence) <= max_chars:
                    current_text += sentence
                else:
                    # セグメント確定
                    current_end = current_start + (len(current_text) * time_per_char)
                    result.append({
                        'start': current_start,
                        'end': min(current_end, end),
                        'text': current_text,
                        'speaker': segment.get('speaker')
                    })
                    current_text = sentence
                    current_start = current_end
            
            # 最後のセグメント
            if current_text:
                result.append({
                    'start': current_start,
                    'end': end,
                    'text': current_text,
                    'speaker': segment.get('speaker')
                })
        
        return result


if __name__ == "__main__":
    # テスト
    logging.basicConfig(level=logging.INFO)
    
    print("=== EnhancedSubtitleExporter Test ===\n")
    
    # テストデータ
    test_segments = [
        {'start': 0.5, 'end': 3.2, 'text': 'こんにちは、今日は会議です。', 'speaker': '話者A'},
        {'start': 3.5, 'end': 6.8, 'text': 'プロジェクトの進捗を確認します。', 'speaker': '話者A'},
        {'start': 7.0, 'end': 10.5, 'text': '昨日の作業は順調でした。', 'speaker': '話者B'},
    ]
    
    exporter = EnhancedSubtitleExporter()
    
    # 各フォーマットテスト
    import tempfile
    import os
    
    with tempfile.TemporaryDirectory() as tmpdir:
        base_path = os.path.join(tmpdir, "test_output")
        
        print("1. Auto Export Test:")
        results = exporter.export_auto(
            test_segments,
            base_path,
            formats=['srt', 'vtt', 'json', 'txt']
        )
        for fmt, success in results.items():
            status = "✓" if success else "✗"
            print(f"   {status} {fmt.upper()}")
        
        print("\n2. SRT Content:")
        srt_content = Path(f"{base_path}.srt").read_text(encoding='utf-8')
        print(srt_content[:200] + "...")
        
        print("\n3. Merge Segments Test:")
        short_segments = [
            {'start': 0.0, 'end': 0.3, 'text': 'あ', 'speaker': 'A'},
            {'start': 0.4, 'end': 0.7, 'text': 'い', 'speaker': 'A'},
            {'start': 0.8, 'end': 3.0, 'text': 'うえお', 'speaker': 'A'},
        ]
        merged = exporter.merge_short_segments(short_segments)
        print(f"   {len(short_segments)} segments -> {len(merged)} segments")
        for seg in merged:
            print(f"     {seg['start']:.1f}s - {seg['end']:.1f}s: {seg['text']}")
    
    print("\nTest completed!")
